from flask import Blueprint, render_template, request, redirect, url_for, flash, session, abort, send_file
import sqlite3
import uuid
from werkzeug.security import generate_password_hash
import os
from docx import Document
from docx.shared import Inches

# Create a Blueprint for the members directory
members_bp = Blueprint('members', __name__, template_folder='templates')

DATABASE = 'church_management.db'

# Helper function to get a database connection
def get_db_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

# Function to log changes in the change_records table
def log_change(user_id, action, target_id=None, target_username=None, change_details=None):
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO change_records (user_id, action, target_id, target_username, change_details)
        VALUES (?, ?, ?, ?, ?)
    ''', (user_id, action, target_id, target_username, change_details))
    conn.commit()
    conn.close()

# Route to display members directory with search functionality
@members_bp.route('/members-directory', methods=['GET', 'POST'])
def members_directory():
    conn = get_db_connection()
    cursor = conn.cursor()

    user_role = session.get('user_role')  # Retrieve the user's role from the session
    current_user_id = session.get('user_id')  # Current logged-in user ID

    # Restrict access to Staff, Admins, and Owners only
    if user_role not in ['Staff', 'Admin', 'Owner']:
        abort(403)  # Forbidden

    search_field = request.args.get('search_field', 'all')
    search_term = request.args.get('search_term', '').strip()

    query = 'SELECT * FROM users'
    params = []

    if search_term:
        if search_field == 'all':
            query += '''
                WHERE first_name LIKE ? OR last_name LIKE ? OR phone LIKE ? OR email LIKE ?
                OR address LIKE ? OR role LIKE ? OR username LIKE ? OR accepts_emails LIKE ?
            '''
            params.extend(['%' + search_term + '%'] * 8)
        else:
            query += f' WHERE {search_field} LIKE ?'
            params.append('%' + search_term + '%')

    cursor.execute(query, params)
    members = cursor.fetchall()

    if request.method == 'POST':
        if user_role == 'Staff':
            flash("You do not have permission to make changes.")
            log_change(user_id=current_user_id, action='unauthorized_attempt',
                       change_details="Staff attempted to edit member")
            return redirect(url_for('members.members_directory'))

        member_id = request.form.get('id')
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        phone = request.form['phone']
        email = request.form['email']
        address = request.form['address']
        username = request.form.get('username', '').strip()  # Default to empty string if not provided
        password = request.form.get('password')  # Will be None if not provided
        role = request.form['role']
        accepts_emails = request.form.get('accepts_emails') == 'Yes'

        # If the username is blank, generate a unique identifier
        if not username:
            username = f"user_{uuid.uuid4().hex[:8]}"  # Generate a unique username using UUID

        # Check for existing username if not a blank auto-generated username
        cursor.execute('SELECT id FROM users WHERE username = ? AND id != ?', (username, member_id or 0))
        existing_user = cursor.fetchone()
        if existing_user:
            flash("The username is already taken. Please choose a different one.")
            return redirect(url_for('members.members_directory'))

        # Role hierarchy checks for Admins
        if user_role == 'Admin':
            if role in ['Admin', 'Owner']:
                flash("You do not have permission to assign Admin or Owner roles.")
                log_change(user_id=current_user_id, action='unauthorized_attempt',
                           change_details="Admin attempted to assign higher role")
                return redirect(url_for('members.members_directory'))

        if password:  # If a password is provided, hash it
            hashed_password = generate_password_hash(password)
        else:  # If no password provided, use None to keep existing password
            hashed_password = None

        if member_id:  # If an ID is present, update the existing member
            cursor.execute('''
                UPDATE users 
                SET first_name = ?, last_name = ?, phone = ?, email = ?, address = ?, username = ?, password = COALESCE(?, password), role = ?, accepts_emails = ?, last_edited_by = ?
                WHERE id = ?
            ''', (
                first_name, last_name, phone, email, address, username, hashed_password, role, accepts_emails,
                current_user_id,
                member_id))
            conn.commit()

            # Log the update action
            log_change(user_id=current_user_id, action='update', target_id=member_id, target_username=username,
                       change_details=f'Updated member: {first_name} {last_name}')

            flash('Member updated successfully.')
        else:  # Otherwise, add a new member
            cursor.execute('''
                INSERT INTO users (first_name, last_name, phone, email, address, username, password, role, accepts_emails, created_by)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                first_name, last_name, phone, email, address, username, hashed_password, role, accepts_emails,
                current_user_id))
            conn.commit()
            new_member_id = cursor.lastrowid

            # Log the creation action
            log_change(user_id=current_user_id, action='create', target_id=new_member_id, target_username=username,
                       change_details=f'Added new member: {first_name} {last_name}')

            flash('Member added successfully.')

        conn.close()
        return redirect(url_for('members.members_directory'))

    # Log the view action
    log_change(user_id=current_user_id, action='view', change_details='Viewed members directory')

    conn.close()
    return render_template('members_directory.html', members=members, user_role=user_role, search_field=search_field, search_term=search_term)

# Route to add a new member
@members_bp.route('/add-member', methods=['GET', 'POST'])
def add_member():
    conn = get_db_connection()
    cursor = conn.cursor()

    user_role = session.get('user_role')  # Retrieve the user's role from the session
    current_user_id = session.get('user_id')  # Current logged-in user ID

    # Restrict access to Admins and Owners only
    if user_role not in ['Admin', 'Owner']:
        abort(403)  # Forbidden

    if request.method == 'POST':
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        phone = request.form['phone']
        email = request.form['email']
        address = request.form['address']
        username = request.form['username']
        password = request.form['password']
        role = request.form['role']
        accepts_emails = request.form.get('accepts_emails') == 'Yes'

        # Role hierarchy checks for Admins
        if user_role == 'Admin':
            if role in ['Admin', 'Owner']:
                flash("You do not have permission to assign Admin or Owner roles.")
                log_change(user_id=current_user_id, action='unauthorized_attempt',
                           change_details="Admin attempted to assign higher role")
                return redirect(url_for('members.add_member'))

        if password:
            hashed_password = generate_password_hash(password)
        else:
            hashed_password = None

        cursor.execute('''
            INSERT INTO users (first_name, last_name, phone, email, address, username, password, role, accepts_emails, created_by)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            first_name, last_name, phone, email, address, username, hashed_password, role, accepts_emails, current_user_id))

        conn.commit()
        new_member_id = cursor.lastrowid

        # Log the creation action
        log_change(user_id=current_user_id, action='create', target_id=new_member_id, target_username=username,
                   change_details=f'Added new member: {first_name} {last_name}')

        conn.close()

        flash('Member added successfully.')
        return redirect(url_for('members.members_directory'))

    conn.close()
    return render_template('add_member.html')

# Route to delete a member
@members_bp.route('/delete-member/<int:id>', methods=['POST'])
def delete_member(id):
    conn = get_db_connection()
    cursor = conn.cursor()

    # Get the role of the member to be deleted
    cursor.execute('SELECT first_name, last_name, username, role FROM users WHERE id = ?', (id,))
    member = cursor.fetchone()
    member_role = member['role']
    member_username = member['username']

    # Get the role of the current user
    user_role = session.get('user_role')
    current_user_id = session.get('user_id')

    # Hierarchy checks
    if user_role == 'Admin' and member_role == 'Owner':
        flash("Admins cannot delete the Owner.")
        log_change(user_id=current_user_id, action='unauthorized_attempt',
                   change_details="Admin attempted to delete Owner without permission")
        return redirect(url_for('members.members_directory'))
    if user_role == 'Staff' and member_role in ['Admin', 'Owner']:
        flash("Staff cannot delete Admins or Owners.")
        log_change(user_id=current_user_id, action='unauthorized_attempt',
                   change_details="Staff attempted to delete Admin or Owner without permission")
        return redirect(url_for('members.members_directory'))
    if current_user_id == id:
        flash("You cannot delete your own account.")
        log_change(user_id=current_user_id, action='unauthorized_attempt',
                   change_details="Attempted to delete own account")
        return redirect(url_for('members.members_directory'))

    cursor.execute('DELETE FROM users WHERE id = ?', (id,))
    conn.commit()

    # Log the deletion action
    log_change(user_id=current_user_id, action='delete', target_id=id, target_username=member_username,
               change_details=f'Deleted member: {member["first_name"]} {member["last_name"]}')

    conn.close()

    flash('Member deleted successfully.')
    return redirect(url_for('members.members_directory'))

# Route to default the members directory to a Word document
@members_bp.route('/default-directory')
def export_directory():
    conn = get_db_connection()
    cursor = conn.cursor()

    user_role = session.get('user_role')  # Retrieve the user's role from the session

    # Restrict access to Admins and Owners only
    if user_role not in ['Admin', 'Owner']:
        abort(403)  # Forbidden

    cursor.execute('SELECT first_name, last_name, phone, email, address, role, username, accepts_emails FROM users')
    members = cursor.fetchall()

    # Create a Word document
    doc = Document()
    doc.add_heading('Members Directory', 0)

    # Set up the table with specific column widths
    table = doc.add_table(rows=1, cols=8)
    table.autofit = False
    col_widths = [Inches(1.0), Inches(1.0), Inches(1.2), Inches(1.5), Inches(2.5), Inches(1.0), Inches(1.0), Inches(1.0)]
    hdr_cells = table.rows[0].cells
    headers = ['First Name', 'Last Name', 'Phone', 'Email', 'Address', 'Role', 'Username', 'Accepts Emails']

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].width = col_widths[i]

    for member in members:
        row_cells = table.add_row().cells
        row_data = [
            member['first_name'], member['last_name'], member['phone'],
            member['email'], member['address'], member['role'],
            member['username'], 'Yes' if member['accepts_emails'] else 'No'
        ]
        for i, data in enumerate(row_data):
            row_cells[i].text = data
            row_cells[i].width = col_widths[i]

    # Save the document to the current working directory
    doc_path = os.path.join(os.getcwd(), 'members_directory.docx')
    doc.save(doc_path)

    conn.close()

    # Send the file to the client
    return send_file(doc_path, as_attachment=True, download_name='members_directory.docx')
