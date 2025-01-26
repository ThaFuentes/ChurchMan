from flask import Blueprint, render_template, redirect, url_for, request, flash, jsonify, session, abort
import sqlite3
from functools import wraps
from docx import Document
import os

# Create a Blueprint for the donations section
donations_bp = Blueprint('donations', __name__, template_folder='templates')

DATABASE = 'church_management.db'


# Helper function to get a database connection
def get_db_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row  # Access columns by name
    return conn


# Function to log changes
def log_change(user_id, action, target_id=None, target_username=None, change_details=None):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO change_records (user_id, action, target_id, target_username, change_details)
        VALUES (?, ?, ?, ?, ?)
    ''', (user_id, action, target_id, target_username, change_details))
    conn.commit()
    conn.close()


# Role required decorator
def role_required(required_roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session or session.get('user_role') not in required_roles:
                abort(403)  # Forbidden
            return f(*args, **kwargs)

        return decorated_function

    return decorator


@donations_bp.route('/')
@role_required(['Staff', 'Admin', 'Owner'])
def donations_dashboard():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    # Fetch data as needed for the dashboard (currently placeholder)
    donations = []  # Replace with actual data fetching logic

    # Log the view action
    log_change(user_id=user_id, action='view', change_details='Viewed donations dashboard')

    return render_template('donations.html', donations=donations)


@donations_bp.route('/add', methods=['GET', 'POST'])
@role_required(['Staff', 'Admin', 'Owner'])
def add_donation():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor()

    # Fetch church information
    cursor.execute('SELECT church_name, address, phone_number, pastor, tax_status FROM church_info LIMIT 1')
    church_info = cursor.fetchone()

    if request.method == 'POST':
        # Process form data and save to the database
        name = request.form['name']
        amount = request.form['amount']
        date = request.form['date']
        method = request.form['method']
        notes = request.form.get('notes', '')

        cursor.execute('''
            INSERT INTO donations (name, amount, date, method, notes)
            VALUES (?, ?, ?, ?, ?)
        ''', (name, amount, date, method, notes))
        conn.commit()

        # Log the addition of the donation
        log_change(user_id=user_id, action='add', change_details=f'Added donation for {name} amount: {amount}')

        conn.close()

        flash('Donation added successfully.')
        return redirect(url_for('donations.donations_dashboard'))

    conn.close()

    return render_template('donation.html', church_info=church_info)


@donations_bp.route('/edit/<int:donation_id>', methods=['GET', 'POST'])
@role_required(['Admin', 'Owner'])
def edit_donation(donation_id):
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor()

    # Fetch the donation record by ID
    cursor.execute('SELECT * FROM donations WHERE id = ?', (donation_id,))
    donation = cursor.fetchone()

    if request.method == 'POST':
        # Update the donation record with the new data
        name = request.form['name']
        amount = request.form['amount']
        date = request.form['date']
        method = request.form['method']
        notes = request.form.get('notes', '')

        cursor.execute('''
            UPDATE donations
            SET name = ?, amount = ?, date = ?, method = ?, notes = ?
            WHERE id = ?
        ''', (name, amount, date, method, notes, donation_id))
        conn.commit()

        # Log the update action
        log_change(user_id=user_id, action='edit', change_details=f'Edited donation for {name} amount: {amount}')

        conn.close()

        flash('Donation updated successfully.')
        return redirect(url_for('donations.donations_dashboard'))

    conn.close()

    return render_template('edit_donation.html', donation=donation)


@donations_bp.route('/delete/<int:donation_id>', methods=['POST'])
@role_required(['Admin', 'Owner'])
def delete_donation(donation_id):
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor()

    # Fetch the donation record to get the name and amount for logging
    cursor.execute('SELECT * FROM donations WHERE id = ?', (donation_id,))
    donation = cursor.fetchone()

    if donation:
        cursor.execute('DELETE FROM donations WHERE id = ?', (donation_id,))
        conn.commit()

        # Log the deletion action
        log_change(user_id=user_id, action='delete',
                   change_details=f'Deleted donation for {donation["name"]} amount: {donation["amount"]}')

        flash('Donation deleted successfully.')

    conn.close()

    return redirect(url_for('donations.donations_dashboard'))


@donations_bp.route('/view_all_members')
@role_required(['Staff', 'Admin', 'Owner'])
def view_all_members():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor()

    # Fetch all members from the database in alphabetical order
    cursor.execute('''
        SELECT id, first_name, last_name, phone, role
        FROM users
        ORDER BY last_name ASC, first_name ASC
    ''')
    members = cursor.fetchall()

    conn.close()

    # Log the view action
    log_change(user_id=user_id, action='view', change_details='Viewed all members')

    return render_template('donations.html', members=members)


@donations_bp.route('/fetch_members_with_donations', methods=['GET'])
@role_required(['Admin', 'Owner'])
def fetch_members_with_donations():
    year = request.args.get('year')
    conn = get_db_connection()
    cursor = conn.cursor()

    if year:
        cursor.execute('''
            SELECT DISTINCT users.id, users.first_name, users.last_name
            FROM users
            INNER JOIN donations ON users.first_name || ' ' || users.last_name = donations.name
            WHERE strftime('%Y', donations.date) = ?
            ORDER BY users.last_name ASC, users.first_name ASC
        ''', (year,))
    else:
        cursor.execute('''
            SELECT DISTINCT users.id, users.first_name, users.last_name
            FROM users
            INNER JOIN donations ON users.first_name || ' ' || users.last_name = donations.name
            ORDER BY users.last_name ASC, users.first_name ASC
        ''')

    members = cursor.fetchall()
    conn.close()

    return jsonify([dict(member) for member in members])


@donations_bp.route('/get_export_location', methods=['GET'])
@role_required(['Admin', 'Owner'])
def get_export_location():
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute('SELECT export_location FROM settings WHERE id = 16')
    result = cursor.fetchone()
    export_location = result['export_location'] if result else ''

    conn.close()

    return jsonify({'export_location': export_location})


@donations_bp.route('/export_individual', methods=['POST'])
@role_required(['Admin', 'Owner'])
def export_individual():
    year = request.form.get('year')
    individual_name = request.form.get('individual_name')
    save_location = request.form.get('save_location')

    conn = get_db_connection()
    cursor = conn.cursor()

    # Fetch church information
    cursor.execute('SELECT church_name, address, phone_number, pastor, tax_status FROM settings LIMIT 1')
    church_info = cursor.fetchone()

    # Fetch the user and their donations
    cursor.execute("SELECT DISTINCT name FROM donations WHERE name = ?", (individual_name,))
    user_name = cursor.fetchone()

    if not user_name:
        flash('Error: No donations found for the specified user.', 'danger')
        return redirect(url_for('donations.donations_dashboard'))

    cursor.execute("SELECT * FROM donations WHERE strftime('%Y', date) = ? AND name = ?", (year, individual_name))
    donations = cursor.fetchall()
    conn.close()

    # Set a default save location if none is provided
    if not save_location:
        save_location = os.path.expanduser('~')  # Default to the user's home directory

    # Construct the full save path
    save_path = os.path.join(save_location, 'Donations', 'Individual', year, f"Donations_{individual_name}.docx")

    try:
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        # Create and save the document
        doc = Document()
        doc.add_heading(f'{church_info["church_name"]}', 0)
        doc.add_paragraph(f'Address: {church_info["address"]}')
        doc.add_paragraph(f'Phone Number: {church_info["phone_number"]}')
        doc.add_paragraph(f'Pastor: {church_info["pastor"]}')
        doc.add_paragraph(f'Tax Status: {church_info["tax_status"]}')
        doc.add_paragraph('')  # Spacer before donations list

        doc.add_heading(f'Individual Donations Report for {individual_name}', level=1)
        for donation in donations:
            doc.add_paragraph(f"Name: {donation['name']}")
            doc.add_paragraph(f"Amount: ${donation['amount']}")
            doc.add_paragraph(f"Date: {donation['date']}")
            doc.add_paragraph(f"Method: {donation['method']}")
            doc.add_paragraph(f"Notes: {donation['notes']}")
            doc.add_paragraph('')

        doc.save(save_path)
        flash(f'Individual donation records exported to {save_path}', 'success')
    except PermissionError as e:
        flash(f'Error saving document: Permission denied. {e}', 'danger')
    except Exception as e:
        flash(f'Error saving document: {str(e)}', 'danger')

    return redirect(url_for('donations.donations_dashboard'))


@donations_bp.route('/export_donations', methods=['GET', 'POST'])
@role_required(['Admin', 'Owner'])
def export_donations():
    if request.method == 'POST':
        export_type = request.form['export_type']
        year = request.form.get('year')
        individual_id = request.form.get('individual_id')
        save_location = request.form.get('save_location')

        conn = get_db_connection()
        cursor = conn.cursor()

        # Fetch donations based on the selected export type
        if export_type == "current_year" or export_type == "total_yearly":
            cursor.execute("SELECT * FROM donations WHERE strftime('%Y', date) = ?", (year,))
        elif export_type == "individual":
            cursor.execute("SELECT * FROM donations WHERE strftime('%Y', date) = ? AND user_id = ?",
                           (year, individual_id))
        elif export_type == "complete":
            cursor.execute("SELECT * FROM donations")

        donations = cursor.fetchall()
        conn.close()

        # Create a Word document
        doc = Document()
        doc.add_heading('Donations Report', 0)
        for donation in donations:
            doc.add_paragraph(f"Name: {donation['name']}")
            doc.add_paragraph(f"Amount: ${donation['amount']}")
            doc.add_paragraph(f"Date: {donation['date']}")
            doc.add_paragraph(f"Method: {donation['method']}")
            doc.add_paragraph(f"Notes: {donation['notes']}")
            doc.add_paragraph('')

        # Determine the save location
        if not save_location:
            # Fetch default save location from settings
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT value FROM settings WHERE key = 'export_location'")
            save_location = cursor.fetchone()['value']
            conn.close()

        if export_type == "individual":
            save_path = os.path.join(save_location, 'Individual', year, f"Donation_{individual_id}.docx")
        elif export_type == "total_yearly":
            save_path = os.path.join(save_location, 'Yearly', year, f"Donation_Year_{year}.docx")
        else:
            save_path = os.path.join(save_location, 'Complete', 'All_Donations.docx')

        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        doc.save(save_path)

        flash(f'Report successfully exported to {save_path}', 'success')
        return redirect(url_for('donations.donations_dashboard'))

    return render_template('export_donations.html')


@donations_bp.route('/reports')
@role_required(['Admin', 'Owner'])
def reports():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    flash("This feature is not yet implemented.")

    # Log the attempt to access the feature
    log_change(user_id=user_id, action='view', change_details='Attempted to view reports')

    return redirect(url_for('donations.donations_dashboard'))


@donations_bp.route('/view_non_member_donations')
@role_required(['Admin', 'Owner'])
def view_non_member_donations():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    flash("This feature is not yet implemented.")

    # Log the attempt to access the feature
    log_change(user_id=user_id, action='view', change_details='Attempted to view non-member donations')

    return redirect(url_for('donations.donations_dashboard'))


@donations_bp.route('/add_user')
@role_required(['Owner'])
def add_user():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    flash("This feature is not yet implemented.")

    # Log the attempt to access the feature
    log_change(user_id=user_id, action='add', change_details='Attempted to add a user')

    return redirect(url_for('donations.donations_dashboard'))


@donations_bp.route('/church_info')
@role_required(['Staff', 'Admin', 'Owner'])
def church_info():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    flash("This feature is not yet implemented.")

    # Log the attempt to access the feature
    log_change(user_id=user_id, action='view', change_details='Attempted to view church information')

    return redirect(url_for('donations.donations_dashboard'))


@donations_bp.route('/view_all_donations', methods=['GET', 'POST'])
@role_required(['Admin', 'Owner'])
def view_all_donations():
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor()

    # Handle search and filter
    search_term = request.args.get('search', '')
    selected_year = request.args.get('year', '')

    if selected_year and selected_year != 'all':
        cursor.execute('''
            SELECT name, SUM(amount) as total_donations, COUNT(*) as number_of_donations
            FROM donations
            WHERE name LIKE ? AND strftime('%Y', date) = ?
            GROUP BY name
            ORDER BY MAX(date) DESC
        ''', ('%' + search_term + '%', selected_year))
    else:
        cursor.execute('''
            SELECT name, SUM(amount) as total_donations, COUNT(*) as number_of_donations
            FROM donations
            WHERE name LIKE ?
            GROUP BY name
            ORDER BY MAX(date) DESC
        ''', ('%' + search_term + '%',))

    donations = cursor.fetchall()

    # Fetch detailed donation data
    detailed_donations = {}
    for donation in donations:
        cursor.execute('''
            SELECT date, amount, method
            FROM donations
            WHERE name = ? AND (? = '' OR strftime('%Y', date) = ?)
            ORDER BY date DESC
        ''', (donation['name'], selected_year, selected_year))
        detailed_donations[donation['name']] = cursor.fetchall()

    # Get distinct years for the dropdown
    cursor.execute('SELECT DISTINCT strftime("%Y", date) as year FROM donations ORDER BY year DESC')
    years = [row['year'] for row in cursor.fetchall()]

    conn.close()

    # Log the view action
    log_change(user_id=user_id, action='view', change_details='Viewed all donations')

    return render_template('view_all_donations.html', donations=donations, detailed_donations=detailed_donations,
                           years=years, selected_year=selected_year)


@donations_bp.route('/get_donation_details/<name>', methods=['GET'])
@role_required(['Admin', 'Owner'])
def get_donation_details(name):
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute('''
        SELECT date, amount, method
        FROM donations
        WHERE name = ?
        ORDER BY date ASC
    ''', (name,))
    donation_details = cursor.fetchall()

    conn.close()

    # Log the view of donation details
    log_change(user_id=user_id, action='view', change_details=f'Viewed donation details for {name}')

    return jsonify([dict(row) for row in donation_details])


@donations_bp.route('/get_years', methods=['GET'])
@role_required(['Admin', 'Owner'])
def get_years():
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute('SELECT DISTINCT strftime("%Y", date) as year FROM donations ORDER BY year DESC')
    years = [row['year'] for row in cursor.fetchall()]

    conn.close()

    return jsonify(years)


@donations_bp.route('/get_donating_members', methods=['GET'])
@role_required(['Admin', 'Owner'])
def get_donating_members():
    conn = get_db_connection()
    cursor = conn.cursor()

    # Fetch distinct years for which there are donations
    cursor.execute('''
        SELECT DISTINCT strftime("%Y", date) as year
        FROM donations
        ORDER BY year DESC
    ''')
    years = cursor.fetchall()

    members_by_year = {}

    for year in years:
        cursor.execute('''
            SELECT u.id, u.first_name, u.last_name, SUM(d.amount) as total_donations
            FROM donations d
            JOIN users u ON d.name = u.first_name || ' ' || u.last_name
            WHERE strftime('%Y', d.date) = ?
            GROUP BY u.id
            ORDER BY total_donations DESC
        ''', (year['year'],))
        members = cursor.fetchall()
        members_by_year[year['year']] = members

    conn.close()

    return jsonify(members_by_year)
